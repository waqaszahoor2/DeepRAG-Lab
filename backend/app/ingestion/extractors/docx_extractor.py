"""DOCX text extractor using python-docx."""

from __future__ import annotations

from dataclasses import dataclass

from app.core.logging import get_logger

logger = get_logger(__name__)


@dataclass
class ExtractedPage:
    page_number: int
    text: str


def extract_docx(file_path: str) -> list[ExtractedPage]:
    """Extract text from a DOCX file.

    DOCX files don't have native page boundaries, so we return the full
    document as a single logical "page" (page_number=1).
    Headings are preserved as section markers.
    """
    from docx import Document

    try:
        doc = Document(file_path)
    except Exception as exc:
        logger.error("DOCX extraction failed for %s: %s", file_path, exc)
        raise

    sections: list[str] = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        # Prefix headings for structure preservation
        if para.style and para.style.name.startswith("Heading"):
            level = para.style.name.replace("Heading", "").strip() or "1"
            sections.append(f"{'#' * int(level)} {text}")
        else:
            sections.append(text)

    # Also extract table content
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                sections.append(row_text)

    full_text = "\n\n".join(sections)

    if not full_text.strip():
        logger.warning("DOCX file appears empty: %s", file_path)
        return []

    logger.info("Extracted %d characters from DOCX: %s", len(full_text), file_path)
    return [ExtractedPage(page_number=1, text=full_text)]
